"""
Credit Appraisal Memo (CAM) Generator.

Generates professional CAM reports in both PDF and DOCX formats.

Structure:
  1. Borrower Overview
  2. Industry Analysis
  3. Financial Analysis
  4. Promoter Background
  5. Risk Analysis
  6. Credit Recommendation
"""
from __future__ import annotations

import os
from datetime import datetime
from pathlib import Path
from typing import Any

from loguru import logger

from config import settings


def _ensure_dir(path: str) -> str:
    os.makedirs(path, exist_ok=True)
    return path


class CAMGenerator:
    """Generates Credit Appraisal Memos in PDF and DOCX formats."""

    OUTPUT_DIR = "./data/reports"

    def __init__(self):
        _ensure_dir(self.OUTPUT_DIR)

    def generate(
        self,
        company_name: str,
        analysis_id: str,
        research_data: dict[str, Any],
        score_data: dict[str, Any],
        financial_signals: dict[str, Any] | None = None,
    ) -> dict[str, str]:
        """
        Generate CAM in both formats.

        Returns:
            {"pdf_path": str, "docx_path": str}
        """
        ctx = self._build_context(company_name, research_data, score_data, financial_signals)
        pdf_path = self._generate_pdf(company_name, analysis_id, ctx)
        docx_path = self._generate_docx(company_name, analysis_id, ctx)
        logger.info(f"[CAM] Generated PDF={pdf_path}, DOCX={docx_path}")
        return {"pdf_path": pdf_path, "docx_path": docx_path}

    # ── Context Builder ───────────────────────────────────────────────────────

    # ── Context Builder ───────────────────────────────────────────────────────

    def _build_context(
        self,
        company_name: str,
        research: dict,
        score: dict,
        signals: dict | None,
    ) -> dict[str, Any]:
        risk_cat = score.get("risk_category", "N/A")
        overall = score.get("overall_credit_score", 0)
        recommendation = self._get_recommendation(risk_cat, overall)

        return {
            "company_name": company_name,
            "date": datetime.now().strftime("%B %d, %Y"),
            "risk_category": risk_cat,
            "overall_score": overall,
            "recommendation": recommendation,
            # Sections
            "borrower_overview": (
                f"{company_name} is the borrowing entity under review. "
                f"The company operates in a sector described as: {research.get('sector_summary', 'N/A')}."
            ),
            "industry_analysis": research.get("sector_summary", "N/A"),
            "financial_analysis": self._format_financial(score, signals),
            "promoter_background": research.get("promoter_summary", "N/A"),
            "risk_analysis": self._format_risk(score, research),
            "credit_recommendation": recommendation,
            "risk_alerts": score.get("risk_alerts", []),
            "indian_intel": score.get("indian_intel", {}),
            "explanation": score.get("explanation", []),
            # Scores
            "character_score": score.get("character_score", 0),
            "capacity_score": score.get("capacity_score", 0),
            "capital_score": score.get("capital_score", 0),
            "collateral_score": score.get("collateral_score", 0),
            "conditions_score": score.get("conditions_score", 0),
            "risk_flags": research.get("risk_flags", []),
        }

    def _format_financial(self, score: dict, signals: dict | None) -> str:
        parts = [f"Overall Credit Score: {score.get('overall_credit_score', 0):.1f}/100"]
        if score.get("score_breakdown", {}).get("ratios"):
            ratios = score["score_breakdown"]["ratios"]
            parts.append(f"  • Debt to Revenue: {ratios.get('debt_to_revenue', 0)*100:.1f}%")
            parts.append(f"  • Current Ratio: {ratios.get('current_ratio', 0):.2f}")
            parts.append(f"  • DSCR: {ratios.get('dscr', 0):.2f}")
        
        if signals:
            for key, val in signals.items():
                if isinstance(val, dict) and val.get("mentioned"):
                    parts.append(f"  • {key.replace('_', ' ').title()}: Detected in documents")
        return "\n".join(parts)

    def _format_risk(self, score: dict, research: dict) -> str:
        flags = research.get("risk_flags", [])
        alerts = score.get("risk_alerts", [])
        lines = [
            f"Character Score: {score.get('character_score', 0):.1f}",
            f"Capacity Score:  {score.get('capacity_score', 0):.1f}",
            f"Capital Score:   {score.get('capital_score', 0):.1f}",
            f"Collateral Score: {score.get('collateral_score', 0):.1f}",
            f"Conditions Score: {score.get('conditions_score', 0):.1f}",
            "",
            "Why This Score (AI Explanation):"
        ]
        lines += [f"  • {e}" for e in score.get("explanation", [])]
        lines.append("")
        
        if alerts:
            lines.append("CRITICAL RISK ALERTS:")
            lines += [f"  ⚠ [{a['type'].upper()}] {a['message']}" for a in alerts]
            lines.append("")

        lines.append("Risk Flags from Research:" if flags else "No major risk flags identified in public records.")
        lines += [f"  ⚠ {f}" for f in flags]
        
        intel = score.get("indian_intel", {})
        if intel:
            lines.append("")
            lines.append("INDIAN FINANCIAL INTELLIGENCE CHECKS:")
            for k, v in intel.items():
                lines.append(f"  • {k.replace('_', ' ').title()}: {v}")
                
        return "\n".join(lines)

    def _get_recommendation(self, risk_cat: str, score: float) -> str:
        if risk_cat == "Low":
            return (
                f"APPROVE — The borrower demonstrates strong creditworthiness with a score of "
                f"{score:.1f}/100 and Low risk profile. Standard covenants apply."
            )
        elif risk_cat == "Medium":
            return (
                f"CONDITIONAL APPROVE — Score of {score:.1f}/100 indicates moderate risk. "
                f"Additional collateral or guarantees recommended. Enhanced monitoring required."
            )
        return (
            f"DECLINE / REFER — Score of {score:.1f}/100 indicates High credit risk. "
            f"Significant red flags present. Refer to senior credit committee for review."
        )

    # ── PDF Generator (ReportLab) ─────────────────────────────────────────────

    def _generate_pdf(self, company_name: str, analysis_id: str, ctx: dict) -> str:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
        )

        filename = f"CAM_{company_name.replace(' ', '_')}_{analysis_id[:8]}.pdf"
        path = str(Path(self.OUTPUT_DIR) / filename)

        doc = SimpleDocTemplate(
            path,
            pagesize=A4,
            rightMargin=2 * cm,
            leftMargin=2 * cm,
            topMargin=2 * cm,
            bottomMargin=2 * cm,
        )
        styles = getSampleStyleSheet()
        primary_color = colors.HexColor("#1e3a5f")
        accent_color = colors.HexColor("#2563eb")

        h1 = ParagraphStyle("H1", parent=styles["Heading1"], textColor=primary_color, fontSize=18, spaceAfter=6)
        h2 = ParagraphStyle("H2", parent=styles["Heading2"], textColor=accent_color, fontSize=13, spaceAfter=4)
        body = ParagraphStyle("Body", parent=styles["Normal"], fontSize=10, leading=14, spaceAfter=6)
        caption = ParagraphStyle("Caption", parent=styles["Normal"], fontSize=8, textColor=colors.grey)

        story = []

        # Title Block
        story.append(Paragraph("CREDIT APPRAISAL MEMORANDUM", h1))
        story.append(Paragraph(f"<b>Borrower:</b> {ctx['company_name']}", body))
        story.append(Paragraph(f"<b>Date:</b> {ctx['date']}", body))
        story.append(HRFlowable(width="100%", thickness=1, color=primary_color))
        story.append(Spacer(1, 0.3 * cm))

        # Score Summary Table
        score_data = [
            ["Five Cs", "Score", "Weight"],
            ["Character", f"{ctx['character_score']:.1f}", "25%"],
            ["Capacity", f"{ctx['capacity_score']:.1f}", "30%"],
            ["Capital", f"{ctx['capital_score']:.1f}", "25%"],
            ["Collateral", f"{ctx['collateral_score']:.1f}", "10%"],
            ["Conditions", f"{ctx['conditions_score']:.1f}", "10%"],
            ["OVERALL", f"{ctx['overall_score']:.1f}", "—"],
        ]
        tbl = Table(score_data, colWidths=[5 * cm, 3 * cm, 2.5 * cm])
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), primary_color),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("BACKGROUND", (0, -1), (-1, -1), accent_color),
            ("TEXTCOLOR", (0, -1), (-1, -1), colors.white),
            ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -2), [colors.white, colors.HexColor("#f0f4ff")]),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("ALIGN", (1, 0), (-1, -1), "CENTER"),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
        ]))
        story.append(tbl)
        story.append(Spacer(1, 0.4 * cm))

        risk_color = {
            "Low": colors.green,
            "Medium": colors.orange,
            "High": colors.red,
        }.get(ctx["risk_category"], colors.grey)

        story.append(Paragraph(
            f'Risk Category: <font color="{risk_color.hexval() if hasattr(risk_color, "hexval") else "black"}">'
            f'<b>{ctx["risk_category"]}</b></font>',
            body,
        ))
        story.append(Spacer(1, 0.3 * cm))

        # Sections
        sections = [
            ("1. Borrower Overview", "borrower_overview"),
            ("2. Industry Analysis", "industry_analysis"),
            ("3. Financial Analysis", "financial_analysis"),
            ("4. Promoter Background", "promoter_background"),
            ("5. Risk Analysis", "risk_analysis"),
            ("6. Credit Recommendation", "credit_recommendation"),
        ]
        for title, key in sections:
            story.append(Paragraph(title, h2))
            content = ctx.get(key, "")
            for line in content.split("\n"):
                if line.strip():
                    story.append(Paragraph(line, body))
            story.append(Spacer(1, 0.3 * cm))

        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.grey))
        story.append(Paragraph(
            "This CAM is generated by CredIntel AI and is intended for internal use only. "
            "All information should be independently verified before credit decisions.",
            caption,
        ))

        doc.build(story)
        return path

    # ── DOCX Generator (python-docx) ──────────────────────────────────────────

    def _generate_docx(self, company_name: str, analysis_id: str, ctx: dict) -> str:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        filename = f"CAM_{company_name.replace(' ', '_')}_{analysis_id[:8]}.docx"
        path = str(Path(self.OUTPUT_DIR) / filename)

        doc = Document()

        # Title
        title = doc.add_heading("CREDIT APPRAISAL MEMORANDUM", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        meta = doc.add_paragraph()
        meta.add_run(f"Borrower: ").bold = True
        meta.add_run(ctx["company_name"])
        meta.add_run(f"\nDate: ").font.bold = True
        meta.add_run(ctx["date"])
        meta.add_run(f"\nRisk Category: ").font.bold = True
        meta.add_run(ctx["risk_category"])
        meta.add_run(f"\nOverall Score: ").font.bold = True
        meta.add_run(f"{ctx['overall_score']:.1f}/100")

        doc.add_heading("Credit Score Summary (Five Cs)", level=1)
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "Dimension"
        hdr[1].text = "Score"
        hdr[2].text = "Weight"

        rows_data = [
            ("Character", ctx["character_score"], "25%"),
            ("Capacity", ctx["capacity_score"], "30%"),
            ("Capital", ctx["capital_score"], "25%"),
            ("Collateral", ctx["collateral_score"], "10%"),
            ("Conditions", ctx["conditions_score"], "10%"),
            ("OVERALL", ctx["overall_score"], "—"),
        ]
        for dim, sc, wt in rows_data:
            row = tbl.add_row().cells
            row[0].text = dim
            row[1].text = f"{sc:.1f}"
            row[2].text = wt

        sections = [
            ("1. Borrower Overview", "borrower_overview"),
            ("2. Industry Analysis", "industry_analysis"),
            ("3. Financial Analysis", "financial_analysis"),
            ("4. Promoter Background", "promoter_background"),
            ("5. Risk Analysis", "risk_analysis"),
            ("6. Credit Recommendation", "credit_recommendation"),
        ]
        for heading, key in sections:
            doc.add_heading(heading, level=1)
            doc.add_paragraph(ctx.get(key, ""))

        doc.add_paragraph(
            "\nThis CAM is generated by CredIntel AI. All information should be "
            "independently verified before credit decisions."
        ).italic = True

        doc.save(path)
        return path
