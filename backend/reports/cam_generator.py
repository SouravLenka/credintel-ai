"""
Credit Appraisal Memo (CAM) Generator.

Generates professional CAM reports in both PDF and DOCX formats.

Structure:
  1. Borrower Overview
  2. Industry Analysis
  3. Financial Analysis
  4. Promoter Background
  5. Risk Analysis
  6. Credit Recommendation
"""
from __future__ import annotations

import os
from datetime import datetime
from pathlib import Path
from typing import Any

from loguru import logger

from config import settings


def _ensure_dir(path: str) -> str:
    os.makedirs(path, exist_ok=True)
    return path


class CAMGenerator:
    """Generates Credit Appraisal Memos in PDF and DOCX formats."""

    OUTPUT_DIR = "./data/reports"

    def __init__(self):
        _ensure_dir(self.OUTPUT_DIR)

    def generate(
        self,
        company_name: str,
        analysis_id: str,
        research_data: dict[str, Any],
        score_data: dict[str, Any],
        financial_signals: dict[str, Any] | None = None,
    ) -> dict[str, str]:
        """
        Generate CAM in both formats.

        Returns:
            {"pdf_path": str, "docx_path": str}
        """
        ctx = self._build_context(company_name, research_data, score_data, financial_signals)
        pdf_path = self._generate_pdf(company_name, analysis_id, ctx)
        docx_path = self._generate_docx(company_name, analysis_id, ctx)
        logger.info(f"[CAM] Generated PDF={pdf_path}, DOCX={docx_path}")
        return {"pdf_path": pdf_path, "docx_path": docx_path}

    # ── Context Builder ───────────────────────────────────────────────────────

    # ── Context Builder ───────────────────────────────────────────────────────

    def _build_context(
        self,
        company_name: str,
        research: dict,
        score: dict,
        signals: dict | None,
    ) -> dict[str, Any]:
        risk_cat = score.get("loan_decision", "N/A")
        overall = score.get("credit_score", 0)
        
        # New weighted breakdown
        breakdown = score.get("score_breakdown", {})
        
        return {
            "company_name": company_name,
            "date": datetime.now().strftime("%B %d, %Y"),
            "risk_category": risk_cat,
            "overall_score": overall,
            "recommendation": self._get_recommendation(risk_cat, overall),
            # Sections
            "borrower_overview": (
                f"{company_name} is the borrowing entity under review. "
                f"The company operates in a sector described as: {research.get('sector_summary', 'N/A')}."
            ),
            "industry_analysis": research.get("sector_summary", "N/A"),
            "financial_analysis": self._format_financial(score, signals),
            "promoter_background": research.get("promoter_summary", "N/A"),
            "risk_analysis": self._format_risk(score, research),
            "credit_recommendation": self._get_recommendation(risk_cat, overall),
            "risk_flags": score.get("risk_flags", []),
            "financial_metrics": score.get("financial_metrics", {}),
            # New Five C weighting attributes
            "financial_strength": breakdown.get("financial_strength", 0),
            "bank_behavior": breakdown.get("bank_behavior", 0),
            "tax_compliance": breakdown.get("tax_compliance", 0),
            "credit_bureau": breakdown.get("credit_bureau", 0),
            "qualitative": breakdown.get("qualitative", 0),
        }

    def _format_financial(self, score: dict, signals: dict | None) -> str:
        parts = [f"Overall Credit Score: {score.get('credit_score', 0):.1f}/100"]
        metrics = score.get("financial_metrics", {})
        if metrics:
            for k, v in metrics.items():
                parts.append(f"  • {k}: {v}")
        
        if signals:
            for key, val in signals.items():
                if isinstance(val, dict) and val.get("mentioned"):
                    parts.append(f"  • {key.replace('_', ' ').title()}: Detected in documents")
        return "\n".join(parts)

    def _format_risk(self, score: dict, research: dict) -> str:
        flags = score.get("risk_flags", [])
        lines = [
            f"Loan Decision: {score.get('loan_decision', 'N/A')}",
            f"Overall Score: {score.get('credit_score', 0):.1f}",
            "",
            "Identified Risk Flags:"
        ]
        if flags:
            lines += [f"  ⚠ {f}" for f in flags]
        else:
            lines.append("  No major risk flags identified.")
                
        return "\n".join(lines)

    def _get_recommendation(self, decision: str, score: float) -> str:
        if decision == "APPROVE":
            return (
                f"APPROVE — The borrower demonstrates strong creditworthiness with a score of "
                f"{score:.1f}/100. Entity meets all automated validation criteria."
            )
        elif decision == "CONDITIONAL APPROVAL":
            return (
                f"CONDITIONAL APPROVE — Score of {score:.1f}/100 indicates moderate risk. "
                f"Enhanced monitoring or additional collateral required."
            )
        elif decision == "MANUAL REVIEW":
            return (
                f"MANUAL REVIEW — Score of {score:.1f}/100. Automated scoring is inconclusive. "
                f"Requires manual credit committee inspection."
            )
        return (
            f"DECLINE / REJECT — The loan application has been automatically REJECTED. "
            f"Score: {score:.1f}/100. Critical validation failures or high risk flags detected."
        )

    # ── PDF Generator (ReportLab) ─────────────────────────────────────────────

    def _generate_pdf(self, company_name: str, analysis_id: str, ctx: dict) -> str:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
        )

        filename = f"CAM_{company_name.replace(' ', '_')}_{analysis_id[:8]}.pdf"
        path = str(Path(self.OUTPUT_DIR) / filename)

        doc = SimpleDocTemplate(
            path,
            pagesize=A4,
            rightMargin=2 * cm,
            leftMargin=2 * cm,
            topMargin=2 * cm,
            bottomMargin=2 * cm,
        )
        styles = getSampleStyleSheet()
        primary_color = colors.HexColor("#1e3a5f")
        accent_color = colors.HexColor("#2563eb")

        h1 = ParagraphStyle("H1", parent=styles["Heading1"], textColor=primary_color, fontSize=18, spaceAfter=6)
        h2 = ParagraphStyle("H2", parent=styles["Heading2"], textColor=accent_color, fontSize=13, spaceAfter=4)
        body = ParagraphStyle("Body", parent=styles["Normal"], fontSize=10, leading=14, spaceAfter=6)
        caption = ParagraphStyle("Caption", parent=styles["Normal"], fontSize=8, textColor=colors.grey)

        story = []

        # Title Block
        story.append(Paragraph("CREDIT APPRAISAL MEMORANDUM", h1))
        story.append(Paragraph(f"<b>Borrower:</b> {ctx['company_name']}", body))
        story.append(Paragraph(f"<b>Date:</b> {ctx['date']}", body))
        story.append(HRFlowable(width="100%", thickness=1, color=primary_color))
        story.append(Spacer(1, 0.3 * cm))

        # Score Summary Table
        score_data = [
            ["Five Cs Factor", "Score", "Max Weight"],
            ["Financial Strength", f"{ctx['financial_strength']:.1f}", "30%"],
            ["Bank behavior", f"{ctx['bank_behavior']:.1f}", "20%"],
            ["Tax compliance", f"{ctx['tax_compliance']:.1f}", "20%"],
            ["Credit Bureau", f"{ctx['credit_bureau']:.1f}", "15%"],
            ["Qualitative Factors", f"{ctx['qualitative']:.1f}", "15%"],
            ["OVERALL SCORE", f"{ctx['overall_score']:.1f}", "100%"],
        ]
        tbl = Table(score_data, colWidths=[5 * cm, 3 * cm, 2.5 * cm])
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), primary_color),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("BACKGROUND", (0, -1), (-1, -1), accent_color),
            ("TEXTCOLOR", (0, -1), (-1, -1), colors.white),
            ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -2), [colors.white, colors.HexColor("#f0f4ff")]),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("ALIGN", (1, 0), (-1, -1), "CENTER"),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
        ]))
        story.append(tbl)
        story.append(Spacer(1, 0.4 * cm))

        risk_color = {
            "APPROVE": colors.green,
            "CONDITIONAL APPROVAL": colors.orange,
            "MANUAL REVIEW": colors.blue,
            "REJECT": colors.red,
        }.get(ctx["risk_category"], colors.grey)

        story.append(Paragraph(
            f'Risk Category: <font color="{risk_color.hexval() if hasattr(risk_color, "hexval") else "black"}">'
            f'<b>{ctx["risk_category"]}</b></font>',
            body,
        ))
        story.append(Spacer(1, 0.3 * cm))

        # Sections
        sections = [
            ("1. Borrower Overview", "borrower_overview"),
            ("2. Industry Analysis", "industry_analysis"),
            ("3. Financial Analysis", "financial_analysis"),
            ("4. Promoter Background", "promoter_background"),
            ("5. Risk Analysis", "risk_analysis"),
            ("6. Credit Recommendation", "credit_recommendation"),
        ]
        for title, key in sections:
            story.append(Paragraph(title, h2))
            content = ctx.get(key, "")
            for line in content.split("\n"):
                if line.strip():
                    story.append(Paragraph(line, body))
            story.append(Spacer(1, 0.3 * cm))

        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.grey))
        story.append(Paragraph(
            "This CAM is generated by CredIntel AI and is intended for internal use only. "
            "All information should be independently verified before credit decisions.",
            caption,
        ))

        doc.build(story)
        return path

    # ── DOCX Generator (python-docx) ──────────────────────────────────────────

    def _generate_docx(self, company_name: str, analysis_id: str, ctx: dict) -> str:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        filename = f"CAM_{company_name.replace(' ', '_')}_{analysis_id[:8]}.docx"
        path = str(Path(self.OUTPUT_DIR) / filename)

        doc = Document()

        # Title
        title = doc.add_heading("CREDIT APPRAISAL MEMORANDUM", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        meta = doc.add_paragraph()
        meta.add_run(f"Borrower: ").bold = True
        meta.add_run(ctx["company_name"])
        meta.add_run(f"\nDate: ").font.bold = True
        meta.add_run(ctx["date"])
        meta.add_run(f"\nRisk Category: ").font.bold = True
        meta.add_run(ctx["risk_category"])
        meta.add_run(f"\nOverall Score: ").font.bold = True
        meta.add_run(f"{ctx['overall_score']:.1f}/100")

        doc.add_heading("Credit Score Summary (Five Cs)", level=1)
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "Dimension"
        hdr[1].text = "Score"
        hdr[2].text = "Weight"

        rows_data = [
            ("Financial Strength", ctx["financial_strength"], "30%"),
            ("Bank behavior", ctx["bank_behavior"], "20%"),
            ("Tax compliance", ctx["tax_compliance"], "20%"),
            ("Credit Bureau", ctx["credit_bureau"], "15%"),
            ("Qualitative", ctx["qualitative"], "15%"),
            ("OVERALL SCORE", ctx["overall_score"], "100%"),
        ]
        for dim, sc, wt in rows_data:
            row = tbl.add_row().cells
            row[0].text = dim
            row[1].text = f"{sc:.1f}"
            row[2].text = wt

        sections = [
            ("1. Borrower Overview", "borrower_overview"),
            ("2. Industry Analysis", "industry_analysis"),
            ("3. Financial Analysis", "financial_analysis"),
            ("4. Promoter Background", "promoter_background"),
            ("5. Risk Analysis", "risk_analysis"),
            ("6. Credit Recommendation", "credit_recommendation"),
        ]
        for heading, key in sections:
            doc.add_heading(heading, level=1)
            doc.add_paragraph(ctx.get(key, ""))

        doc.add_paragraph(
            "\nThis CAM is generated by CredIntel AI. All information should be "
            "independently verified before credit decisions."
        ).italic = True

        doc.save(path)
        return path
